"""
Generate Word Document Report for BU SMARTNESS & Alignment Analysis
Shows averaged (symmetric) alignment values
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_symmetric_alignment_matrix(excel_file):
    """Create symmetric alignment matrix by averaging both directions"""
    excel_data = pd.ExcelFile(excel_file)
    bu_names = [sheet for sheet in excel_data.sheet_names if sheet != 'Summary']
    
    alignment_matrix = pd.DataFrame(index=bu_names, columns=bu_names, dtype=float)
    
    for bu1 in bu_names:
        df1 = pd.read_excel(excel_file, sheet_name=bu1)
        
        for bu2 in bu_names:
            if bu1 == bu2:
                alignment_matrix.loc[bu1, bu2] = 100.0
            else:
                # Get alignment from bu1 to bu2
                alignment_section1 = df1[
                    df1.iloc[:, 0].notna() & 
                    ~df1.iloc[:, 0].astype(str).str.contains('Goal ID|ALIGNMENT', case=False, na=False)
                ]
                align_1_to_2 = None
                for _, row in alignment_section1.iterrows():
                    if row.iloc[0] == bu2:
                        align_1_to_2 = pd.to_numeric(row.iloc[1], errors='coerce')
                        break
                
                # Get alignment from bu2 to bu1
                df2 = pd.read_excel(excel_file, sheet_name=bu2)
                alignment_section2 = df2[
                    df2.iloc[:, 0].notna() & 
                    ~df2.iloc[:, 0].astype(str).str.contains('Goal ID|ALIGNMENT', case=False, na=False)
                ]
                align_2_to_1 = None
                for _, row in alignment_section2.iterrows():
                    if row.iloc[0] == bu1:
                        align_2_to_1 = pd.to_numeric(row.iloc[1], errors='coerce')
                        break
                
                # Average both directions
                if align_1_to_2 is not None and align_2_to_1 is not None:
                    alignment_matrix.loc[bu1, bu2] = (align_1_to_2 + align_2_to_1) / 2
                elif align_1_to_2 is not None:
                    alignment_matrix.loc[bu1, bu2] = align_1_to_2
                elif align_2_to_1 is not None:
                    alignment_matrix.loc[bu1, bu2] = align_2_to_1
                else:
                    alignment_matrix.loc[bu1, bu2] = 0.0
    
    return alignment_matrix

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_bu_alignment_section(doc, bu_name, alignment_matrix):
    """Add alignment section for a specific BU"""
    # BU heading
    heading = doc.add_heading(f"{bu_name}", level=2)
    heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    
    # Get alignment data for this BU
    bu_alignments = []
    for other_bu in alignment_matrix.columns:
        if other_bu != bu_name:
            align_val = alignment_matrix.loc[bu_name, other_bu]
            if pd.notna(align_val):
                bu_alignments.append({
                    'Business Unit': other_bu,
                    'Alignment %': align_val
                })
    
    # Sort by alignment percentage
    bu_alignments_df = pd.DataFrame(bu_alignments).sort_values('Alignment %', ascending=False)
    
    # Create table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Business Unit'
    header_cells[1].text = 'Mutual Alignment %'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for _, row in bu_alignments_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Business Unit']
        row_cells[1].text = f"{row['Alignment %']:.1f}%"
    
    doc.add_paragraph()  # Add spacing

def generate_word_report(excel_file, output_file):
    """Generate comprehensive Word report"""
    print("Loading data...")
    summary_df = pd.read_excel(excel_file, sheet_name='Summary')
    alignment_matrix = create_symmetric_alignment_matrix(excel_file)
    
    print("Creating Word document...")
    doc = Document()
    
    # Title
    title = doc.add_heading('BU SMARTNESS & Alignment Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(108, 117, 125)
    
    doc.add_paragraph()
    
    # Executive Summary
    add_heading_with_style(doc, '1. Executive Summary', level=1)
    
    total_bus = len(summary_df)
    total_goals = summary_df['Total Goals'].sum()
    avg_smartness = summary_df['Average SMARTNESS %'].mean()
    
    summary_text = doc.add_paragraph()
    summary_text.add_run(f"• Total Business Units: ").bold = True
    summary_text.add_run(f"{total_bus}\n")
    summary_text.add_run(f"• Total Goals Analyzed: ").bold = True
    summary_text.add_run(f"{total_goals}\n")
    summary_text.add_run(f"• Average SMARTNESS Score: ").bold = True
    summary_text.add_run(f"{avg_smartness:.1f}%\n")
    
    doc.add_paragraph()
    
    # SMARTNESS Scores Section
    add_heading_with_style(doc, '2. SMARTNESS Scores by Business Unit', level=1)
    
    doc.add_paragraph(
        "SMARTNESS scores measure how well each goal meets the SMART criteria "
        "(Specific, Measurable, Achievable, Relevant, Time-bound). Scores are calculated "
        "using Azure OpenAI semantic analysis."
    )
    
    # SMARTNESS table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Business Unit'
    header_cells[1].text = 'Average SMARTNESS %'
    header_cells[2].text = 'Total Goals'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Business Unit']
        row_cells[1].text = f"{row['Average SMARTNESS %']:.1f}%"
        row_cells[2].text = str(row['Total Goals'])
    
    doc.add_page_break()
    
    # Alignment Analysis Section
    add_heading_with_style(doc, '3. Cross-BU Alignment Analysis', level=1)
    
    doc.add_paragraph(
        "Alignment percentages show mutual alignment between Business Units. "
        "Values are calculated by averaging bidirectional semantic similarity: "
        "how well BU A's goals align with BU B's objectives and vice versa. "
        "This provides a symmetric measure of cross-BU collaboration and goal alignment."
    )
    
    doc.add_paragraph()
    
    # Add alignment section for each BU
    for bu_name in alignment_matrix.index:
        add_bu_alignment_section(doc, bu_name, alignment_matrix)
    
    # Methodology Section
    doc.add_page_break()
    add_heading_with_style(doc, '4. Methodology', level=1)
    
    methodology = doc.add_paragraph()
    methodology.add_run("SMARTNESS Calculation:\n").bold = True
    methodology.add_run(
        "Each goal is analyzed using Azure OpenAI (GPT-4o) to evaluate how well it meets "
        "the SMART criteria. The model assigns a percentage score based on the presence and "
        "quality of Specific, Measurable, Achievable, Relevant, and Time-bound elements.\n\n"
    )
    
    methodology.add_run("Alignment Calculation:\n").bold = True
    methodology.add_run(
        "Cross-BU alignment is calculated using semantic similarity analysis. For each BU pair:\n"
        "  1. Goals from BU A are compared against objectives of BU B\n"
        "  2. Goals from BU B are compared against objectives of BU A\n"
        "  3. The two directional scores are averaged to produce a symmetric mutual alignment percentage\n\n"
    )
    
    methodology.add_run("Data Source:\n").bold = True
    methodology.add_run(
        "All BU objectives are stored in goals.db database. User goals are stored in db.sqlite3. "
        "Analysis was performed using Azure OpenAI endpoint with rate limiting to ensure quality results."
    )
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Note: This report contains aggregated metrics only. Individual goal content is not displayed to maintain privacy.")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = RGBColor(108, 117, 125)
    
    # Save document
    print(f"Saving report to {output_file}...")
    doc.save(output_file)
    print(f"✓ Report generated successfully: {output_file}")

if __name__ == "__main__":
    excel_file = "BU_SMARTNESS_ALIGNMENT_REPORT_20260507_171137.xlsx"
    output_file = f"BU_SMARTNESS_ALIGNMENT_REPORT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    generate_word_report(excel_file, output_file)
